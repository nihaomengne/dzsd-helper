from flask import Flask, render_template, request, send_file, redirect, session, flash, url_for, abort
from docx import Document
from docx.shared import Mm, Pt
from io import BytesIO
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from PIL import Image
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.middleware.proxy_fix import ProxyFix
import sqlite3
import openpyxl
import urllib.parse
import os
import functools
import secrets


def env_bool(name: str, default: bool = False) -> bool:
    value = os.environ.get(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.environ.get('DATA_DIR', os.path.join(BASE_DIR, 'data'))
os.makedirs(DATA_DIR, exist_ok=True)
DB_PATH = os.path.join(DATA_DIR, 'document_counter.db')
DEFAULT_MAX_COUNT = int(os.environ.get('DEFAULT_MAX_COUNT', '10'))
ALLOW_PUBLIC_REGISTRATION = env_bool('ALLOW_PUBLIC_REGISTRATION', False)
INIT_ADMIN_USERNAME = os.environ.get('INIT_ADMIN_USERNAME', 'admin')
INIT_ADMIN_PASSWORD = os.environ.get('INIT_ADMIN_PASSWORD')
TRUST_PROXY = env_bool('TRUST_PROXY', True)

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY') or secrets.token_hex(32)
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE='Lax',
    SESSION_COOKIE_SECURE=env_bool('SESSION_COOKIE_SECURE', True),
    MAX_CONTENT_LENGTH=int(os.environ.get('MAX_CONTENT_LENGTH_MB', '20')) * 1024 * 1024,
)

if TRUST_PROXY:
    app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)


@app.after_request
def add_security_headers(response):
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    response.headers['Cache-Control'] = 'no-store'
    response.headers['Content-Security-Policy'] = (
        "default-src 'self'; "
        "img-src 'self' data: https:; "
        "style-src 'self' 'unsafe-inline' https://cdnjs.cloudflare.com; "
        "script-src 'self' 'unsafe-inline' https://cdnjs.cloudflare.com; "
        "font-src 'self' https://cdnjs.cloudflare.com; "
        "connect-src 'self'; "
        "frame-ancestors 'self';"
    )
    return response


def get_db_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def fit_image_to_a4(img_width, img_height):
    a4_width, a4_height = 210, 297
    width_ratio = a4_width / img_width
    height_ratio = a4_height / img_height
    if width_ratio < height_ratio:
        return a4_width, img_height * width_ratio
    return img_width * height_ratio, a4_height


def init_db():
    conn = get_db_connection()
    c = conn.cursor()

    c.execute('''CREATE TABLE IF NOT EXISTS users
                (id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                is_admin INTEGER DEFAULT 0)''')

    c.execute('''CREATE TABLE IF NOT EXISTS user_document_counter
                (user_id INTEGER UNIQUE,
                count INTEGER DEFAULT 0,
                max_count INTEGER DEFAULT 10,
                FOREIGN KEY(user_id) REFERENCES users(id))''')

    c.execute('''CREATE TABLE IF NOT EXISTS DocumentCounter (count INTEGER)''')
    c.execute('SELECT COUNT(*) FROM DocumentCounter')
    if c.fetchone()[0] == 0:
        c.execute('INSERT INTO DocumentCounter (count) VALUES (0)')

    c.execute("PRAGMA table_info(users)")
    user_columns = {row[1] for row in c.fetchall()}
    if 'password' in user_columns and 'password_hash' not in user_columns:
        c.execute('ALTER TABLE users ADD COLUMN password_hash TEXT')
        c.execute('SELECT id, password FROM users WHERE password IS NOT NULL')
        for row in c.fetchall():
            c.execute('UPDATE users SET password_hash = ? WHERE id = ?',
                      (generate_password_hash(row['password']), row['id']))

    c.execute('SELECT id, username, password_hash FROM users')
    for user in c.fetchall():
        if not user['password_hash']:
            fallback_password = os.environ.get('MIGRATION_FALLBACK_PASSWORD', 'ChangeMeNow123!')
            c.execute('UPDATE users SET password_hash = ? WHERE id = ?',
                      (generate_password_hash(fallback_password), user['id']))

    c.execute('SELECT COUNT(*) FROM users WHERE is_admin = 1')
    admin_exists = c.fetchone()[0] > 0
    if not admin_exists:
        if not INIT_ADMIN_PASSWORD:
            raise RuntimeError('缺少 INIT_ADMIN_PASSWORD 环境变量，无法初始化管理员账号')
        c.execute('INSERT INTO users (username, password_hash, is_admin) VALUES (?, ?, ?)',
                  (INIT_ADMIN_USERNAME, generate_password_hash(INIT_ADMIN_PASSWORD), 1))
        user_id = c.lastrowid
        c.execute('INSERT OR REPLACE INTO user_document_counter (user_id, count, max_count) VALUES (?, ?, ?)',
                  (user_id, 0, 999999))

    c.execute('SELECT id FROM users')
    for row in c.fetchall():
        c.execute('INSERT OR IGNORE INTO user_document_counter (user_id, count, max_count) VALUES (?, 0, ?)',
                  (row['id'], DEFAULT_MAX_COUNT))

    conn.commit()
    conn.close()


def current_user_is_admin():
    return bool(session.get('is_admin'))


def login_required(f):
    @functools.wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('请先登录', 'error')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function


def admin_required(f):
    @functools.wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('请先登录', 'error')
            return redirect(url_for('login'))
        if not current_user_is_admin():
            flash('需要管理员权限', 'error')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username'].strip()
        password = request.form['password']

        conn = get_db_connection()
        c = conn.cursor()
        c.execute('SELECT id, username, password_hash, is_admin FROM users WHERE username = ?', (username,))
        user = c.fetchone()
        conn.close()

        if user and check_password_hash(user['password_hash'], password):
            session.clear()
            session['user_id'] = user['id']
            session['username'] = user['username']
            session['is_admin'] = bool(user['is_admin'])
            return redirect(url_for('index'))

        flash('用户名或密码错误', 'error')

    return render_template('login.html', allow_public_registration=ALLOW_PUBLIC_REGISTRATION)


@app.route('/logout')
def logout():
    session.clear()
    flash('已退出登录', 'info')
    return redirect(url_for('login'))


@app.route('/admin', methods=['GET', 'POST'])
@admin_required
def admin_panel():
    conn = get_db_connection()
    c = conn.cursor()

    if request.method == 'POST':
        if 'add_user' in request.form:
            username = request.form['username'].strip()
            password = request.form['password']
            max_count = int(request.form['max_count'])
            is_admin = 1 if 'is_admin' in request.form else 0

            if len(password) < 8:
                flash('密码至少 8 位', 'error')
            else:
                try:
                    c.execute('INSERT INTO users (username, password_hash, is_admin) VALUES (?, ?, ?)',
                              (username, generate_password_hash(password), is_admin))
                    user_id = c.lastrowid
                    c.execute('INSERT INTO user_document_counter (user_id, count, max_count) VALUES (?, ?, ?)',
                              (user_id, 0, max_count))
                    conn.commit()
                    flash(f'用户 {username} 已添加', 'success')
                except sqlite3.IntegrityError:
                    flash(f'用户名 {username} 已存在', 'error')

        elif 'update_count' in request.form:
            user_id = int(request.form['user_id'])
            max_count = int(request.form['new_max_count'])
            c.execute('UPDATE user_document_counter SET max_count = ? WHERE user_id = ?', (max_count, user_id))
            conn.commit()
            flash('用户文档生成次数已更新', 'success')

    c.execute('''SELECT u.id, u.username, u.is_admin, udc.count, udc.max_count
                FROM users u
                JOIN user_document_counter udc ON u.id = udc.user_id
                ORDER BY u.id ASC''')
    users = c.fetchall()
    conn.close()
    return render_template('admin.html', users=users)


@app.route('/', methods=['GET', 'POST'])
@login_required
def index():
    conn = get_db_connection()
    c = conn.cursor()

    c.execute('SELECT count FROM DocumentCounter')
    total_count = c.fetchone()[0]

    c.execute('SELECT count, max_count FROM user_document_counter WHERE user_id = ?', (session['user_id'],))
    user_count_data = c.fetchone()
    user_count = user_count_data['count']
    max_count = user_count_data['max_count']
    is_admin = current_user_is_admin()
    reached_limit = user_count >= max_count and not is_admin

    if request.method == 'POST':
        if reached_limit:
            conn.close()
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return {'status': 'error', 'message': '您已达到文档生成次数上限，请联系管理员增加次数'}, 403
            flash('您已达到文档生成次数上限，请联系管理员增加次数', 'error')
            return redirect(url_for('index'))

        branch = request.form['branch']
        application_date = request.form['application_date']
        merchant_name = request.form['merchant_name']
        merchant_address = request.form['merchant_address']
        merchant_code = request.form['merchant_code']
        contact_person = request.form['contact_person']
        contact_phone = request.form['contact_phone']
        operator = request.form['operator']
        bank_card_number = request.form.get('bank_card_number', '')

        formatted_date = datetime.strptime(application_date, '%Y-%m-%d').strftime('%Y年%m月%d日')

        if 'search_qcc' in request.form:
            encoded_name = urllib.parse.quote(merchant_name)
            return redirect(f'https://www.qcc.com/web/search?key={encoded_name}')

        doc = Document(os.path.join(BASE_DIR, 'word_template.docx'))
        new_text = f"申请支行（签章）： {branch}   申请日期：{formatted_date}"

        for paragraph in doc.paragraphs:
            if "申请支行（签章）： 保定银行定州支行" in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run(new_text)
                run.font.name = '仿宋'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                run.font.size = Pt(12)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            if "商户：" in paragraph.text:
                parts = paragraph.text.split("商户：")
                if len(parts) > 1:
                    paragraph.clear()
                    paragraph.add_run(parts[0] + "商户：")
                    run = paragraph.add_run(merchant_name)
                    run.font.name = '仿宋'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                    if "日期：" in parts[1]:
                        paragraph.add_run("   " + parts[1][parts[1].find("日期："):])
            if "商户签章：" in paragraph.text and paragraph.text.endswith("商户签章："):
                paragraph.add_run(merchant_name)
            if "承诺商户：" in paragraph.text and paragraph.text.endswith("承诺商户："):
                paragraph.add_run(merchant_name)

        tables = doc.tables
        for i in range(2):
            tables[i].cell(1, 1).text = merchant_name
            tables[i].cell(2, 1).text = merchant_address
            tables[i].cell(9, 1).text = merchant_address
            tables[i].cell(3, 1).text = merchant_code
            tables[i].cell(10, 1).text = contact_person
            tables[i].cell(10, 5).text = contact_phone
            tables[i].cell(11, 1).text = merchant_name
            tables[i].cell(12, 1).text = operator

        for i in range(2, 6):
            tables[i].cell(0, 1).text = merchant_name
            tables[i].cell(0, 3).text = branch

        for paragraph in doc.paragraphs:
            if "培训人员：" in paragraph.text:
                paragraph.text = paragraph.text.replace("培训人员：", f"培训人员：{operator}")
            if "巡检人员：" in paragraph.text:
                paragraph.text = paragraph.text.replace("巡检人员：", f"巡检人员：{operator}")
            if "卡号：" in paragraph.text:
                paragraph.text = paragraph.text.replace("卡号：", f"卡号：{bank_card_number}")

        for i in range(6, 10):
            for row in tables[i].rows:
                for cell in row.cells:
                    if "商户名称" in cell.text:
                        cell.text = f"商户名称：{merchant_name}"
                    elif "商户代码" in cell.text:
                        cell.text = f"商户代码：{merchant_code}"
                    elif "商户地址" in cell.text:
                        cell.text = f"商户地址：{merchant_address}"
                    elif "所属支行" in cell.text:
                        cell.text = f"所属支行：{branch}"

        images = request.files.getlist('images')
        if images and images[0].filename:
            for image in images:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                section.page_height = Mm(297)
                section.page_width = Mm(210)
                section.left_margin = Mm(0)
                section.right_margin = Mm(0)
                section.top_margin = Mm(0)
                section.bottom_margin = Mm(0)
                with Image.open(image) as img:
                    img_width, img_height = img.size
                image.stream.seek(0)
                new_width, new_height = fit_image_to_a4(img_width, img_height)
                doc.add_picture(image, width=Mm(new_width), height=Mm(new_height))

        if 'generate_explanation' in request.form:
            doc.add_section(WD_SECTION.NEW_PAGE)
            title = doc.add_paragraph()
            title_run = title.add_run("情况说明")
            title_run.font.name = '宋体'
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            title_run.font.size = Pt(18)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()
            content = doc.add_paragraph()
            content.paragraph_format.first_line_indent = Pt(28)
            content_text = f'我支行营销办理的"{merchant_name}"收单业务，商户门头地址与营业执照相符，但名称不符。经客户经理核实，该店面所在位置与营业执照地址相统一，经营者相关信息真实完整，经营范围符合实际情况，确认为同一商户。'
            content_run = content.add_run(content_text)
            content_run.font.name = '仿宋'
            content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            content_run.font.size = Pt(16)
            explain = doc.add_paragraph()
            explain.paragraph_format.first_line_indent = Pt(28)
            explain_run = explain.add_run("特此说明。")
            explain_run.font.name = '仿宋'
            explain_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            explain_run.font.size = Pt(16)
            for _ in range(3):
                doc.add_paragraph()
            footer = doc.add_paragraph()
            footer_run = footer.add_run(f"保定银行股份有限公司{branch}")
            footer_run.font.name = '仿宋'
            footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            footer_run.font.size = Pt(16)
            footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            date = doc.add_paragraph()
            date_run = date.add_run(formatted_date)
            date_run.font.name = '仿宋'
            date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            date_run.font.size = Pt(16)
            date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        if 'generate_excel' in request.form:
            wb = openpyxl.load_workbook(os.path.join(BASE_DIR, 'excel_template.xlsx'))
            ws = wb.active
            ws['D2'] = branch
            ws['B9'] = ws['D9'] = application_date.replace("年", "  ").replace("月", "  ").replace("日", "  ").replace("-", "")
            ws['B2'] = merchant_name
            ws['B4'] = ws['B22'] = contact_person
            ws['B6'] = ws['B20'] = merchant_address
            ws['D4'] = operator
            ws['B23'] = bank_card_number
            ws['B24'] = f"保定银行股份有限公司{branch}"
            excel_output = BytesIO()
            wb.save(excel_output)
            excel_output.seek(0)
            return send_file(excel_output, as_attachment=True, download_name=f'{merchant_name}.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        c.execute('UPDATE DocumentCounter SET count = count + 1')
        c.execute('UPDATE user_document_counter SET count = count + 1 WHERE user_id = ?', (session['user_id'],))
        conn.commit()
        conn.close()
        return send_file(output, as_attachment=True, download_name=f'{merchant_name}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    conn.close()
    return render_template(
        'index.html',
        count=total_count,
        user_count=user_count,
        max_count=max_count,
        username=session.get('username'),
        reached_limit=reached_limit,
        today=datetime.now().strftime('%Y-%m-%d')
    )


@app.route('/change_password', methods=['GET', 'POST'])
@login_required
def change_password():
    if request.method == 'POST':
        old_password = request.form['old_password']
        new_password = request.form['new_password']
        confirm_password = request.form['confirm_password']

        if new_password != confirm_password:
            flash('新密码不匹配', 'error')
            return redirect(url_for('change_password'))
        if len(new_password) < 8:
            flash('新密码至少 8 位', 'error')
            return redirect(url_for('change_password'))

        conn = get_db_connection()
        c = conn.cursor()
        c.execute('SELECT password_hash FROM users WHERE id = ?', (session['user_id'],))
        user = c.fetchone()

        if user and check_password_hash(user['password_hash'], old_password):
            c.execute('UPDATE users SET password_hash = ? WHERE id = ?',
                      (generate_password_hash(new_password), session['user_id']))
            conn.commit()
            conn.close()
            flash('密码修改成功', 'success')
            return redirect(url_for('index'))

        conn.close()
        flash('原密码错误', 'error')
        return redirect(url_for('change_password'))

    return render_template('change_password.html')


@app.route('/admin/reset_password/<int:user_id>', methods=['POST'])
@admin_required
def admin_reset_password(user_id):
    new_password = request.form.get('new_password', '')
    if len(new_password) < 8:
        flash('新密码至少 8 位', 'error')
        return redirect(url_for('admin_panel'))

    conn = get_db_connection()
    c = conn.cursor()
    try:
        c.execute('UPDATE users SET password_hash = ? WHERE id = ?',
                  (generate_password_hash(new_password), user_id))
        conn.commit()
        flash('密码重置成功', 'success')
    except Exception as e:
        flash(f'密码重置失败: {str(e)}', 'error')
    finally:
        conn.close()
    return redirect(url_for('admin_panel'))


@app.route('/show_register')
def show_register():
    if not ALLOW_PUBLIC_REGISTRATION:
        abort(404)
    return render_template('register.html')


@app.route('/register', methods=['GET', 'POST'])
def register():
    if not ALLOW_PUBLIC_REGISTRATION:
        abort(404)
    if request.method == 'GET':
        return render_template('register.html')

    username = request.form['username'].strip()
    password = request.form['password']
    confirm_password = request.form['confirm_password']

    if password != confirm_password:
        flash('两次输入的密码不一致', 'error')
        return render_template('register.html', username=username)
    if len(password) < 8:
        flash('密码至少 8 位', 'error')
        return render_template('register.html', username=username)

    conn = get_db_connection()
    c = conn.cursor()
    try:
        c.execute('INSERT INTO users (username, password_hash, is_admin) VALUES (?, ?, 0)',
                  (username, generate_password_hash(password)))
        user_id = c.lastrowid
        c.execute('INSERT INTO user_document_counter (user_id, count, max_count) VALUES (?, 0, 0)', (user_id,))
        conn.commit()
        flash('注册成功，请联系管理员开通次数', 'success')
        return redirect(url_for('login', username=username))
    except sqlite3.IntegrityError:
        flash('用户名已存在', 'error')
        return render_template('register.html', username=username)
    finally:
        conn.close()


if __name__ == '__main__':
    init_db()
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', '1012')), debug=False)
else:
    init_db()
